"""
Data Loader cho Personal AI Agent — Knowledge Base.

Module này chịu trách nhiệm load và preprocess dữ liệu nguồn từ thư mục data/
để đưa vào pipeline RAG (embed → store → retrieve).

Kiến trúc: Strategy Pattern + Registry
    - Mỗi loại file (CSV, MD, PDF, DOCX) có một loader class riêng
    - Tất cả loader được tự động đăng ký vào LoaderRegistry khi import module
    - DataLoader là facade chính, tự detect file type và gọi đúng loader

Cách mở rộng khi thêm file type mới:
    1. Tạo class NewLoader(BaseFileLoader)
    2. Implement load() method và set supported_extensions
    3. Thêm auto-register ở cuối section ĐĂNG KÝ CÁC LOADER MẶC ĐỊNH
    → Done! DataLoader sẽ tự nhận dạng file type mới

Các loader hiện có:
    - CSVLoader:      .csv
    - MarkdownLoader: .md
    - PDFLoader:      .pdf  (sử dụng PyMuPDF)
    - DOCXLoader:     .docx (sử dụng python-docx)

Cách sử dụng:
    from knowledge_base.documents_loader import DataLoader

    loader = DataLoader()

    # Load một file cụ thể
    docs = loader.load_file("report.pdf")

    # Load toàn bộ thư mục
    all_docs = loader.load_all()

    # Kiểm tra data có thay đổi không (dùng cho indexer)
    current_hash = loader.compute_hash("report.pdf")
"""

from __future__ import annotations

import csv
import hashlib
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import ClassVar

from core.exceptions import DataLoadError
from core.logger import get_logger

logger = get_logger(__name__)
CHUNK_SIZE=1200 
CHUNK_OVERLAP=200


# ═══════════════════════════════════════════════════════════════════════
# DATA MODEL — Cấu trúc dữ liệu chuẩn cho mọi loại document
# ═══════════════════════════════════════════════════════════════════════

@dataclass
class Document:
    """
    Đơn vị dữ liệu chuẩn sau khi load, dùng cho toàn bộ pipeline downstream.

    Attributes:
        doc_id: ID duy nhất cho document (dùng làm ChromaDB document ID).
        content: Nội dung text chính — sẽ được embed và lưu vào vector store.
        metadata: Thông tin bổ sung (source file, class, question gốc, ...).
    """
    doc_id: str
    content: str
    metadata: dict = field(default_factory=dict)


# ═══════════════════════════════════════════════════════════════════════
# BASE LOADER — Interface chung cho mọi file type
# ═══════════════════════════════════════════════════════════════════════

class BaseFileLoader(ABC):
    """
    Abstract base class cho tất cả file loader.

    Mỗi subclass cần implement:
        - load(file_path) → list[Document]: logic đọc và parse file
        - supported_extensions: danh sách extension hỗ trợ (ví dụ: [".csv"])
    """

    supported_extensions: ClassVar[list[str]] = []

    @abstractmethod
    def load(self, file_path: Path) -> list[Document]:
        """
        Load và parse file thành danh sách Document.

        Args:
            file_path: Đường dẫn tuyệt đối hoặc tương đối tới file.

        Returns:
            Danh sách Document đã được preprocess.

        Raises:
            DataLoadError: Khi file không tồn tại, format sai, hoặc rỗng.
        """
        pass

    def _validate_file_exists(self, file_path: Path) -> None:
        """Kiểm tra file có tồn tại không."""
        if not file_path.exists():
            raise DataLoadError(
                f"File not found: {file_path}",
                details={"file": str(file_path)},
            )
        if not file_path.is_file():
            raise DataLoadError(
                f"Path is not a file: {file_path}",
                details={"file": str(file_path)},
            )


# ═══════════════════════════════════════════════════════════════════════
# CSV LOADER — Load file CSV (BankFAQs, branch_info, ...)
# ═══════════════════════════════════════════════════════════════════════

class CSVLoader(BaseFileLoader):
    """
    Loader cho file CSV.

    Hỗ trợ 2 chế độ:
        1. combine_columns: Nối nhiều cột thành một text (cho FAQ — cần embed)
        2. raw: Mỗi row thành 1 Document với toàn bộ columns trong metadata
           (cho branch_info — dùng trực tiếp, không cần embed)

    Args:
        combine_columns: List tên cột sẽ được nối thành content.
                         Nếu None, sử dụng toàn bộ các cột.
        separator: Ký tự phân tách khi nối các cột.
        required_columns: List tên cột bắt buộc phải có. Raise lỗi nếu thiếu.
        id_prefix: Prefix cho doc_id (ví dụ: "faq", "branch").
        encoding: Encoding của file CSV.
    """

    supported_extensions: ClassVar[list[str]] = [".csv"]

    def __init__(
        self,
        combine_columns: list[str] | None = None,
        separator: str = " | ",
        required_columns: list[str] | None = None,
        id_prefix: str = "doc",
        encoding: str = "utf-8-sig",
        chunk_size: int = CHUNK_SIZE,
        chunk_overlap: int = CHUNK_OVERLAP,
    ):
        self.combine_columns = combine_columns
        self.separator = separator
        self.required_columns = required_columns or []
        self.id_prefix = id_prefix
        self.encoding = encoding
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def load(self, file_path: Path) -> list[Document]:
        """
        Load file CSV và chuyển thành danh sách Document.

        Quy trình:
            1. Validate file tồn tại
            2. Đọc CSV với DictReader
            3. Validate required columns
            4. Với mỗi row: tạo content (combine hoặc raw) + metadata
            5. Bỏ qua row rỗng, log warning
            6. Validate kết quả không rỗng

        Returns:
            list[Document]: Danh sách document đã preprocess.

        Raises:
            DataLoadError: File không tồn tại, thiếu cột, hoặc data rỗng.
        """
        file_path = Path(file_path)
        self._validate_file_exists(file_path)

        logger.info(f"Loading CSV: {file_path.name}")

        try:
            from langchain_text_splitters import RecursiveCharacterTextSplitter
        except ImportError as e:
            raise DataLoadError(
                "langchain-text-splitters is required to load and chunk CSV files. "
                "Install it with: pip install langchain-text-splitters",
                details={"file": str(file_path), "error": str(e)},
            )

        try:
            with open(file_path, mode="r", encoding=self.encoding, newline="") as f:
                reader = csv.DictReader(f)

                # --- Validate header ---
                if reader.fieldnames is None:
                    raise DataLoadError(
                        "CSV file has no header row",
                        details={"file": str(file_path)},
                    )

                actual_columns = list(reader.fieldnames)
                self._validate_columns(actual_columns, file_path)

                # --- Xác định columns dùng cho content ---
                content_columns = self.combine_columns or actual_columns

                # --- Khởi tạo splitter ---
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=self.chunk_size,
                    chunk_overlap=self.chunk_overlap,
                    length_function=len,
                    is_separator_regex=False,
                )

                # --- Parse rows ---
                documents: list[Document] = []
                skipped = 0

                for row_idx, row in enumerate(reader):
                    # Tạo content bằng cách nối các cột
                    parts = []
                    for col in content_columns:
                        value = (row.get(col) or "").strip()
                        if value:
                            parts.append(value)

                    row_content = self.separator.join(parts)

                    # Bỏ qua row rỗng
                    if not row_content.strip():
                        skipped += 1
                        continue

                    # Chia chunk cho content của row này
                    chunks = text_splitter.split_text(row_content)

                    for chunk_idx, chunk in enumerate(chunks):
                        content = chunk.strip()
                        if not content:
                            continue

                        # Dùng content-hash thay vì sequential index
                        hash_input = f"{file_path.name}::row_{row_idx}::chunk_{chunk_idx}::{content}"
                        content_hash = hashlib.md5(hash_input.encode("utf-8")).hexdigest()[:12]
                        doc_id = f"{self.id_prefix}_{content_hash}"
                        metadata = {
                            "source_file": file_path.name,
                            "row_index": row_idx,
                            "chunk_index": chunk_idx,
                            **{k: (v or "").strip() for k, v in row.items()},
                        }
    
                        documents.append(Document(
                            doc_id=doc_id,
                            content=content,
                            metadata=metadata,
                        ))

                # --- Log kết quả ---
                if skipped > 0:
                    logger.warning(
                        f"Skipped {skipped} empty rows in {file_path.name}"
                    )

                if not documents:
                    raise DataLoadError(
                        "CSV file produced no documents after preprocessing",
                        details={
                            "file": str(file_path),
                            "skipped": skipped,
                        },
                    )

                logger.info(
                    f"Loaded {len(documents)} documents from {file_path.name}"
                )
                return documents

        except DataLoadError:
            # Re-raise DataLoadError — không wrap lại
            raise
        except UnicodeDecodeError as e:
            raise DataLoadError(
                f"Encoding error reading {file_path.name}",
                details={
                    "file": str(file_path),
                    "encoding": self.encoding,
                    "error": str(e),
                },
            ) from e
        except Exception as e:
            raise DataLoadError(
                f"Unexpected error loading {file_path.name}",
                details={"file": str(file_path), "error": str(e)},
            ) from e

    def _validate_columns(self, actual_columns: list[str], file_path: Path) -> None:
        """Kiểm tra file CSV có đủ các cột bắt buộc không."""
        if not self.required_columns:
            return

        missing = [col for col in self.required_columns if col not in actual_columns]
        if missing:
            raise DataLoadError(
                f"CSV file is missing required columns: {missing}",
                details={
                    "file": str(file_path),
                    "required": self.required_columns,
                    "actual": actual_columns,
                    "missing": missing,
                },
            )

# ═══════════════════════════════════════════════════════════════════════
# MD LOADER — Load file MD
# ═══════════════════════════════════════════════════════════════════════

class MarkdownLoader(BaseFileLoader):
    """Loader cho file Markdown (.md)."""

    supported_extensions: ClassVar[list[str]] = [".md"]

    def __init__(self, chunk_separator: str = "\n## ", id_prefix: str = "md"):
        self.chunk_separator = chunk_separator
        self.id_prefix = id_prefix

    def load(self, file_path: Path) -> list[Document]:
        file_path = Path(file_path)
        self._validate_file_exists(file_path)

        logger.info(f"Loading Markdown: {file_path.name}")

        text = file_path.read_text(encoding="utf-8")

        # Tách theo heading ## 
        chunks = text.split(self.chunk_separator)
        
        documents = []
        for idx, chunk in enumerate(chunks):
            content = chunk.strip()
            if not content:
                continue

            # Dùng content-hash thay vì sequential index
            hash_input = f"{file_path.name}::chunk_{idx}::{content}"
            content_hash = hashlib.md5(hash_input.encode("utf-8")).hexdigest()[:12]
            doc_id = f"{self.id_prefix}_{content_hash}"

            documents.append(Document(
                doc_id=doc_id,
                content=content,
                metadata={
                    "source_file": file_path.name,
                    "chunk_index": idx,
                },
            ))

        logger.info(f"Loaded {len(documents)} documents from {file_path.name}")
        return documents

# ═══════════════════════════════════════════════════════════════════════
# PDF LOADER — Load file PDF (sử dụng PyMuPDF)
# ═══════════════════════════════════════════════════════════════════════

class PDFLoader(BaseFileLoader):
    """
    Loader cho file PDF sử dụng PyMuPDF (fitz) và RecursiveCharacterTextSplitter.

    Trích xuất text từ từng trang PDF, sau đó sử dụng text splitter
    để chia nhỏ nội dung thành các chunk có kích thước cố định và overlap.

    Args:
        id_prefix: Prefix cho doc_id (mặc định: "pdf").
        chunk_size: Kích thước tối đa của mỗi chunk (mặc định: 1000 ký tự).
        chunk_overlap: Số ký tự gối lên nhau giữa các chunk (mặc định: 200).
    """

    supported_extensions: ClassVar[list[str]] = [".pdf"]

    def __init__(self, id_prefix: str = "pdf", chunk_size: int = 1000, chunk_overlap: int = 200):
        self.id_prefix = id_prefix
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def load(self, file_path: Path) -> list[Document]:
        file_path = Path(file_path)
        self._validate_file_exists(file_path)

        logger.info(f"Loading PDF: {file_path.name} with chunk_size={self.chunk_size}")

        try:
            import fitz  # PyMuPDF
            from langchain_text_splitters import RecursiveCharacterTextSplitter
        except ImportError as e:
            raise DataLoadError(
                "PyMuPDF and langchain-text-splitters are required to load PDF files. "
                "Install them with: pip install PyMuPDF langchain-text-splitters",
                details={"file": str(file_path), "error": str(e)},
            )

        try:
            doc = fitz.open(file_path)
            documents: list[Document] = []
            
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                length_function=len,
                is_separator_regex=False,
            )

            full_text = ""
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text().strip()
                if text:
                    full_text += text + "\n\n"

            # Chia text của toàn bộ file thành các chunk
            chunks = text_splitter.split_text(full_text)
            
            for idx, chunk in enumerate(chunks):
                content = chunk.strip()
                if not content:
                    continue
                    
                # Dùng content-hash cho doc_id
                hash_input = f"{file_path.name}::chunk_{idx}::{content}"
                content_hash = hashlib.md5(hash_input.encode("utf-8")).hexdigest()[:12]
                doc_id = f"{self.id_prefix}_{content_hash}"

                documents.append(Document(
                    doc_id=doc_id,
                    content=content,
                    metadata={
                        "source_file": file_path.name,
                        "total_pages": len(doc),
                        "chunk_index": idx,
                    },
                ))

            doc.close()

            if not documents:
                raise DataLoadError(
                    "PDF file produced no documents (no extractable text)",
                    details={"file": str(file_path)},
                )

            logger.info(
                f"Loaded {len(documents)} chunks from {file_path.name}"
            )
            return documents

        except DataLoadError:
            raise
        except Exception as e:
            raise DataLoadError(
                f"Error reading PDF file: {file_path.name}",
                details={"file": str(file_path), "error": str(e)},
            ) from e


# ═══════════════════════════════════════════════════════════════════════
# DOCX LOADER — Load file DOCX (sử dụng python-docx)
# ═══════════════════════════════════════════════════════════════════════

class DOCXLoader(BaseFileLoader):
    """
    Loader cho file Word DOCX sử dụng python-docx.

    Trích xuất text từ tất cả paragraph trong file DOCX.
    Gộp toàn bộ paragraph thành một Document duy nhất cho mỗi file.

    Args:
        id_prefix: Prefix cho doc_id (mặc định: "docx").
    """

    supported_extensions: ClassVar[list[str]] = [".docx"]

    def __init__(self, id_prefix: str = "docx", chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP):
        self.id_prefix = id_prefix
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def load(self, file_path: Path) -> list[Document]:
        file_path = Path(file_path)
        self._validate_file_exists(file_path)

        logger.info(f"Loading DOCX: {file_path.name}")

        try:
            from docx import Document as DocxDocument
            from langchain_text_splitters import RecursiveCharacterTextSplitter
        except ImportError as e:
            raise DataLoadError(
                "python-docx and langchain-text-splitters are required. "
                "Install them with: pip install python-docx langchain-text-splitters",
                details={"file": str(file_path), "error": str(e)},
            )

        try:
            docx_doc = DocxDocument(file_path)

            # Trích xuất text từ tất cả paragraph
            paragraphs = [
                para.text.strip()
                for para in docx_doc.paragraphs
                if para.text.strip()
            ]

            if not paragraphs:
                raise DataLoadError(
                    "DOCX file produced no documents (no text content)",
                    details={"file": str(file_path)},
                )

            full_text = "\n\n".join(paragraphs)

            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                length_function=len,
                is_separator_regex=False,
            )

            chunks = text_splitter.split_text(full_text)
            
            documents = []
            for idx, chunk in enumerate(chunks):
                content = chunk.strip()
                if not content:
                    continue
                    
                # Dùng content-hash cho doc_id
                hash_input = f"{file_path.name}::chunk_{idx}::{content}"
                content_hash = hashlib.md5(hash_input.encode("utf-8")).hexdigest()[:12]
                doc_id = f"{self.id_prefix}_{content_hash}"

                documents.append(Document(
                    doc_id=doc_id,
                    content=content,
                    metadata={
                        "source_file": file_path.name,
                        "total_paragraphs": len(paragraphs),
                        "chunk_index": idx,
                        
                    },
                ))

            logger.info(
                f"Loaded {len(documents)} chunks from {file_path.name} "
                f"({len(paragraphs)} paragraphs)"
            )
            return documents

        except DataLoadError:
            raise
        except Exception as e:
            raise DataLoadError(
                f"Error reading DOCX file: {file_path.name}",
                details={"file": str(file_path), "error": str(e)},
            ) from e




# ═══════════════════════════════════════════════════════════════════════
# LOADER REGISTRY — Đăng ký và quản lý loader theo file extension
# ═══════════════════════════════════════════════════════════════════════

class LoaderRegistry:
    """
    Registry trung tâm quản lý mapping: file extension → loader instance.

    Cho phép đăng ký loader mới mà không sửa code DataLoader.

    Cách dùng:
        # Đăng ký loader mới
        LoaderRegistry.register(".md", MarkdownLoader())

        # Lấy loader cho file type
        loader = LoaderRegistry.get_loader(".csv")

        # Xem tất cả extension được hỗ trợ
        exts = LoaderRegistry.supported_extensions()
    """

    _registry: ClassVar[dict[str, BaseFileLoader]] = {}

    @classmethod
    def register(cls, extension: str, loader: BaseFileLoader) -> None:
        """
        Đăng ký một loader cho extension cụ thể.

        Args:
            extension: File extension (bao gồm dấu chấm, ví dụ: ".csv", ".md").
            loader: Instance của BaseFileLoader subclass.
        """
        ext = extension.lower()
        cls._registry[ext] = loader
        logger.debug(f"Registered loader for '{ext}': {loader.__class__.__name__}")

    @classmethod
    def get_loader(cls, extension: str) -> BaseFileLoader | None:
        """
        Lấy loader tương ứng với extension.

        Returns:
            BaseFileLoader nếu có, None nếu chưa đăng ký.
        """
        return cls._registry.get(extension.lower())

    @classmethod
    def supported_extensions(cls) -> list[str]:
        """Trả về danh sách extension đã đăng ký."""
        return list(cls._registry.keys())

    @classmethod
    def is_supported(cls, extension: str) -> bool:
        """Kiểm tra extension có được hỗ trợ không."""
        return extension.lower() in cls._registry

    @classmethod
    def clear(cls) -> None:
        """Xóa toàn bộ registry (dùng cho testing)."""
        cls._registry.clear()


# ═══════════════════════════════════════════════════════════════════════
# ĐĂNG KÝ CÁC LOADER MẶC ĐỊNH
# Tự động đăng ký tất cả loader khi module được import.
# ═══════════════════════════════════════════════════════════════════════

def _auto_register_loaders() -> None:
    """
    Tự động đăng ký tất cả loader đã định nghĩa vào LoaderRegistry.

    Mỗi loader class khai báo `supported_extensions` — hàm này duyệt qua
    tất cả loader và đăng ký từng extension tương ứng.
    """
    default_loaders: list[BaseFileLoader] = [
        CSVLoader(),
        MarkdownLoader(),
        PDFLoader(),
        DOCXLoader(),
    ]

    for loader in default_loaders:
        for ext in loader.supported_extensions:
            LoaderRegistry.register(ext, loader)


# Chạy auto-register khi module được import
_auto_register_loaders()



# ═══════════════════════════════════════════════════════════════════════
# DATA LOADER — Facade chính cho toàn bộ hệ thống
# ═══════════════════════════════════════════════════════════════════════

class DataLoader:
    """
    Facade chính để load dữ liệu từ thư mục raw.

    Tự động detect file type, chọn đúng loader, và trả về list[Document].
    Hỗ trợ:
        - Load một file cụ thể
        - Load toàn bộ thư mục
        - Compute content hash để detect data changes (cho indexer)

    Ví dụ:
        loader = DataLoader(data_dir="data/raw")
        all_docs = loader.load_all()
        faq_docs = loader.load_file("BankFAQs.csv")
    """

    def __init__(self, data_dir: str | Path = "data/raw"):
        """
        Args:
            data_dir: Đường dẫn tới thư mục chứa raw data.
        """
        self.data_dir = Path(data_dir)

    def load_file(self, filename: str) -> list[Document]:
        """
        Load một file cụ thể từ data_dir.

        Tự động chọn loader phù hợp dựa trên file extension
        thông qua LoaderRegistry.

        Args:
            filename: Tên file (ví dụ: "report.pdf", "notes.md").

        Returns:
            list[Document]: Danh sách document từ file.

        Raises:
            DataLoadError: File không tồn tại hoặc extension chưa được hỗ trợ.
        """
        file_path = self.data_dir / filename

        # Chọn loader theo extension từ registry
        ext = file_path.suffix.lower()
        loader = LoaderRegistry.get_loader(ext)

        if loader is None:
            supported = LoaderRegistry.supported_extensions()
            raise DataLoadError(
                f"No loader registered for file type: {file_path.suffix}",
                details={
                    "file": filename,
                    "extension": file_path.suffix,
                    "supported_extensions": supported,
                },
            )

        return loader.load(file_path)

    def load_all(self) -> list[Document]:
        """
        Load toàn bộ file được hỗ trợ trong data_dir.

        Quét tất cả file trong thư mục, bỏ qua file có extension chưa đăng ký.

        Returns:
            list[Document]: Tất cả document từ mọi file.

        Raises:
            DataLoadError: Thư mục data_dir không tồn tại.
        """
        if not self.data_dir.exists():
            raise DataLoadError(
                f"Data directory not found: {self.data_dir}",
                details={"data_dir": str(self.data_dir)},
            )

        if not self.data_dir.is_dir():
            raise DataLoadError(
                f"Path is not a directory: {self.data_dir}",
                details={"data_dir": str(self.data_dir)},
            )

        all_documents: list[Document] = []
        loaded_files: list[str] = []
        skipped_files: list[str] = []

        for file_path in sorted(self.data_dir.iterdir()):
            if not file_path.is_file():
                continue

            filename = file_path.name

            if not LoaderRegistry.is_supported(file_path.suffix.lower()):
                skipped_files.append(filename)
                continue

            try:
                docs = self.load_file(filename)
                all_documents.extend(docs)
                loaded_files.append(filename)
            except DataLoadError as e:
                logger.error(f"Failed to load {filename}: {e}")
                # Tiếp tục load các file khác, không dừng toàn bộ
                continue

        logger.info(
            f"Directory scan complete: "
            f"loaded {len(loaded_files)} files ({len(all_documents)} docs), "
            f"skipped {len(skipped_files)} unsupported files"
        )

        if skipped_files:
            logger.debug(f"Skipped files: {skipped_files}")

        return all_documents

    def compute_hash(self, filename: str) -> str:
        """
        Tính MD5 hash của file content — dùng cho indexer detect data changes.

        Args:
            filename: Tên file trong data_dir.

        Returns:
            MD5 hex digest string.

        Raises:
            DataLoadError: File không tồn tại.
        """
        file_path = self.data_dir / filename

        if not file_path.exists():
            raise DataLoadError(
                f"Cannot compute hash — file not found: {file_path}",
                details={"file": str(file_path)},
            )

        hasher = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                hasher.update(chunk)

        hash_value = hasher.hexdigest()
        logger.debug(f"Hash for {filename}: {hash_value}")
        return hash_value

    def compute_directory_hash(self) -> str:
        """
        Tính combined hash cho toàn bộ thư mục raw data.
        Dùng để kiểm tra nhanh xem có file nào thay đổi không.

        Returns:
            MD5 hex digest string.
        """
        combined_hasher = hashlib.md5()

        for file_path in sorted(self.data_dir.iterdir()):
            if not file_path.is_file():
                continue

            # Bao gồm tên file trong hash (detect rename/delete)
            combined_hasher.update(file_path.name.encode("utf-8"))

            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(8192), b""):
                    combined_hasher.update(chunk)

        hash_value = combined_hasher.hexdigest()
        logger.debug(f"Directory hash for {self.data_dir}: {hash_value}")
        return hash_value
